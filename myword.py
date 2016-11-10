#
# Markdown to Word with scientific helpers (captions,cite)
#
# Emanuele Ruffaldi 2016

# TODO emit images
# TODO emit cite
# TODO emit ref
# TODO emit footnote
# TODO emit caption
# TODO more numbered p1.set_numbering(numbering['ilvl'], self.numId)
# TODO nested lists: https://github.com/python-openxml/python-docx/issues/122
import mistune
import sys
import docx
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm
from docxhelper import add_hyperlink,para_set_numbering
# process text if any
# emit link around text if any
class Link:
    def __init__(self,link,title,text,auto):
        self.link = link
        self.title = title
        if text is None or text == "":
            text = link
        self.text = text
        self.auto = auto

# process text and then apply style
#p.add_run('bold').bold = True
#p.add_run('italic.').italic = True
class Style:
    def __init__(self,text,mode):
        self.text = text
        self.mode = mode

# add image needs download
# add caption
# add reference
#document.add_picture('monty-truth.png', width=Inches(1.25))
class Image:
    def __init__(self,src,title,text):
        self.src = src
        self.title = title
        self.text = text

# process and add footnote
class Footnote:
    def __init__(self,content):
        self.text = content

# bookmark WHERE? TODO
class Bookmark:
    def __init__(self,label):
        self.label = label

# plain text
class Run:
    def __init__(self,text):
        self.text = text

# add REF tag + bibtex
class Cite:
    def __init__(self,cite):
        self.cite = cite

# add bookmark
class Caption:
    def __init__(self,ref):
        self.ref = ref

# add REF tag
class Ref:
    def __init__(self,ref):
        self.ref = ref

class Holder:
    def __init__(self):
        self.runs = []
    def __iadd__(self,r):
        if type(r) is list:
            self.runs.extend(r)
        elif type(r) == Holder:
            self.runs.extend(r.runs)
        else:
            if type(r) is str:
                r = Run(r)
            self.runs.append(r)
        return self
    def __len__(self):
        return len(self.runs)

#p = document.add_paragraph('A plain paragraph having some ')
class Paragraph(Holder):
    def __init__(self,ho):
        Holder.__init__(self)
        self.runs = ho.runs
        # Analize for Images
        if len(self.runs) >= 3:
            print [x.__class__ for x in self.runs]
            # Link == the image
            # Run dummy
            # Caption 
            # rest
            if isinstance(ho.runs[0],Link) and isinstance(ho.runs[1],Run) and isinstance(ho.runs[2],Caption):
                img = Image(ho.runs[0].link,ho.runs[2].ref,ho.runs[3])
                print "replace image ignore:",ho.runs[1].text,"with img"
                self.runs = [img]


class List(Holder):
    def __init__(self,ho,ord):
        Holder.__init__(self)
        self.runs = ho.runs
        self.ordered = ord

class ListItem(Holder):
    def __init__(self,ho):
        Holder.__init__(self)
        self.runs = ho.runs

#document.add_heading('Document Title', 0)
#document.add_heading('Heading, level 1', level=1)
class Header:
    def __init__(self,text,level,raw=None):
        self.text = text
        self.level = level
        self.raw = raw

class Document(Holder):
    def __init__(self,r):
        Holder.__init__(self)
        self.runs = r.runs

class WordRenderer(mistune.Renderer):
    def block_code(self, code, lang):
        print "*code",len(code),lang
        return ""
    def placeholder(self):
        return Holder()
    def header(self, text, level,raw=None):
        print "*head",text,level
        return Header(text,level,raw)
    def link(self, link, title, text):
        print "*link",link,"title=",title,"text=",text
        # TODO download image and cache
        return Link(link,title,text,False)
    def autolink(self, link, is_email=False):
        print "*alink",link
        # TODO detect image and cache (if start of line)
        return Link(link,"","",True)
    def image(self, src, title, text):
        print "*image",title,src,text
        # TODO download image and cache
        return Image(src,title,text)
    def paragraph(self, text):
        # prints the RENDERED
        print "*paragraph",len(text),text
        return Paragraph(text)
    def list(self, body, ordered=True):
        # native
        return List(body,ordered)
    def list_item(self, text):
        # rendered
        return ListItem(text)
    def double_emphasis(self, text):
        return Style(text,"bold")
    def emphasis(self, text):    
        return Style(text,"italics")
    def text(self, text):
        # native
        # FIX \cite{}
        # FIX \ref{}
        # AUTO detect caption
        print "*text",text
        w = [text]
        # special with recursion
        pieces = []

        # order matters, footnote last
        pieces = [("\\footnote{",Footnote),("\\cite{",Cite),("\\ref{",Ref),("\\caption{",Caption),("\\label{",Bookmark),("\\url{",lambda x: Link(x,None,None,None))]
        for p,t in pieces:
            r = []
            print "proecessing",w
            for text in w:
                if type(text) is str:
                    a = text.split(p)
                    if len(a) == 1:
                        r.append(a[0])
                    else:
                        if len(a[0]) != 0:
                            r.append(a[0])
                        for i in range(1,len(a)):
                            # text cite}text cite}text 
                            if t == Footnote:
                                # look for correct number of parent
                                n = 1
                                aw = a[i]
                                j = 0
                                for j in range(0,len(aw)):
                                    if aw[j] == '{':
                                        n += 1
                                    elif aw[j] == '}':
                                        n -= 1
                                        if n == 0:
                                            break
                                inside = aw[0:j]
                                outside = aw[j+1:]
                                r.append(t(self.text(inside)))
                                if len(outside) != 0:
                                    r.append(self.text(outside))
                            else:
                                aw = a[i].split("}",1)
                                r.append(t(aw[0]))
                                if len(aw) == 2:
                                    r.append(Run(aw[1]))
                else:
                    r.append(text)
            w = r
        return [Run(x) if type(x) is str else x for x in w]

class Wordizer:
    def __init__(self,d):
        self.doc = d
    def run(self,ctx,x):
        n = x.__class__.__name__
        m = getattr(self,"do_" + n,None)
        if m is not None:
            return m(ctx,x)
        else:
            print "skip",n
    def gettext(self,q):
        if isinstance(q,Holder):
            s = ""
            for p in q.runs:
                s += self.gettext(p)
            return s
        elif isinstance(q,Run):
            return q.text
        else:
            print "gettext",q.__class__.__name__
            return q.text
    def do_Document(self,ctx,x):
        for p in x.runs:
            self.run(self.doc,p)
    def do_Header(self,doc,x):
        return doc.add_heading(self.gettext(x.text),level=x.level)
    def do_Run(self,up,x):
        return up.add_run(x.text)
    def do_Paragraph(self,up,x):
        para = up.add_paragraph()
        for p in x.runs:
            self.run(para,p)
        return para        
    def do_Link(self,up,x):
        return add_hyperlink(up,x.link,x.text)
    def do_Holder(self,up,x):
        q = []
        for p in x.runs:
            q.append(self.run(up,p))
        return q
    def do_Ref(self,up,x):
        return up.add_run("<<"+x.ref+">>")
    def do_Cite(self,up,x):
        return up.add_run(x.cite)
    def do_ListItem(self,upx,x):
        up,style = upx
        para = self.doc.add_paragraph()
        #para_set_numbering(para,2,2)
        para.style = style
        for p in x.runs:
            self.run(para,p)
        return para
    def do_Style(self,up,x):
        rr = self.run(up,x.text)
        for r in rr:
            if x.mode == "bold":
                r.bold = True
            elif x.mode == "italics":
                print "apply ita to",r.text
                r.italic = True
        return rr
    def do_List(self,up,x):
        s = "List Number" if x.ordered else "List Bullet"
        for p in x.runs:
            r = self.run((up,s),p)
            # if x is ordered ...
        return None

#miss list paragraph 
def dump(r,w=""):
    if isinstance(r,Holder):
        print w,"holder",r
        w += " "
        for k in r.runs:
            dump(k,w)
    else:
        print w,r
def main():
    renderer = WordRenderer()
    markdown = mistune.Markdown(renderer=renderer)  
    r = Document(markdown(open(sys.argv[1],"rb").read()))
    dump(r)
    document = docx.Document()
    w = Wordizer(document)
    w.run(document,r)
    document.save("x.docx")
    print "generated x.docx"



if __name__ == '__main__':
    main()
